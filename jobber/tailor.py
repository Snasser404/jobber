"""Generate a tailored cover letter + lightly-tuned resume for one job."""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

import docx
from docx.shared import Pt

from .models import RankedJob, TailoredApplication
from . import llm, config

SYSTEM = """You are a professional resume writer helping a REAL candidate apply to a specific job.
Be strictly HONEST: only re-emphasize and rephrase the candidate's real experience to fit the job.
Never invent employers, titles, dates, metrics, degrees, or skills the candidate does not already have.

Candidate's master resume:
-----
{resume}
-----
Candidate: {name} | {email} | {location}"""

PROMPT = """Create tailored application materials for this job.

JOB TITLE: {title}
COMPANY: {company}
LOCATION: {location}
JOB DESCRIPTION:
{description}

Skills/keywords to weave in naturally — ONLY where they are genuinely true for the candidate:
{keywords}

Return ONLY a JSON object:
{{
  "summary": "<a rewritten 'About Me' summary, 3-4 sentences, first person, tuned to this role, using only true facts from the resume>",
  "cover_letter": "<a complete, specific cover letter, 250-350 words, addressed to the hiring team at {company}. Reference the role and 2-3 concrete achievements drawn from the resume. Warm and professional, not generic. Do NOT use bracket placeholders.>",
  "resume_changes": ["<plain-language description of each change you made and why>", "..."],
  "emphasize_skills": ["<existing skill to surface for this job>", "..."]
}}"""


# Resume section headers whose following paragraph holds the summary we tailor.
SUMMARY_HEADERS = {
    "about me", "about", "summary", "professional summary", "profile",
    "professional profile", "objective", "career summary", "career objective",
    "personal statement", "profile summary",
}


def _safe(text: str, n: int = 40) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "_", text or "").strip("_")
    return s[:n] or "x"


def _set_text(paragraph, text: str) -> None:
    """Replace a paragraph's text while keeping its existing formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def _build_resume(master_path: str, summary: str, out_path: Path) -> Path:
    """Copy the master resume and swap in the tailored 'About Me' summary."""
    doc = docx.Document(str(config.abspath(master_path)))
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if p.text.strip().rstrip(":").lower() in SUMMARY_HEADERS:
            for j in range(i + 1, len(paras)):
                if paras[j].text.strip():
                    if summary:
                        _set_text(paras[j], summary)
                    break
            break
    doc.save(str(out_path))
    return out_path


def _build_cover_letter(text: str, applicant: dict, out_path: Path) -> Path:
    doc = docx.Document()
    head = doc.add_paragraph()
    run = head.add_run(applicant.get("full_name", ""))
    run.bold = True
    run.font.size = Pt(14)
    contact = " | ".join(x for x in [
        applicant.get("email", ""), applicant.get("phone", ""),
        applicant.get("location", ""), applicant.get("linkedin", ""),
    ] if x)
    if contact:
        doc.add_paragraph(contact)
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph("")
    for block in text.split("\n\n"):
        if block.strip():
            doc.add_paragraph(block.strip())
    doc.save(str(out_path))
    return out_path


def generate(ranked: RankedJob, profile: dict, cfg: dict) -> TailoredApplication:
    job = ranked.job
    applicant = profile.get("applicant", {})
    system = SYSTEM.format(
        resume=profile.get("resume_text", ""),
        name=applicant.get("full_name", ""),
        email=applicant.get("email", ""),
        location=applicant.get("location", ""),
    )
    data = llm.complete_json(
        PROMPT.format(
            title=job.title, company=job.company, location=job.location,
            description=job.description[:5000],
            keywords=", ".join(ranked.keywords) or "(none specified)",
        ),
        system=system, role="write",
        max_tokens=2000, temperature=0.5, cache_system=True,
    )
    summary = str(data.get("summary", "")).strip()
    cover = str(data.get("cover_letter", "")).strip()
    changes = list(data.get("resume_changes", []) or [])

    out_dir = config.abspath(cfg.get("output_dir", "data/output"))
    out_dir.mkdir(parents=True, exist_ok=True)
    base = f"{_safe(job.company)}__{_safe(job.title)}"

    resume_path = _build_resume(profile["resume_path"], summary, out_dir / f"{base}__resume.docx")
    cover_path = _build_cover_letter(cover, applicant, out_dir / f"{base}__cover_letter.docx")

    return TailoredApplication(
        job_id=job.id,
        cover_letter=cover,
        resume_changes=changes,
        resume_docx_path=str(resume_path),
        cover_letter_path=str(cover_path),
        tailored_summary=summary,
    )
