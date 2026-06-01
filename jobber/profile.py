"""Read the user's master resume + preferences into a profile the AI uses."""
from __future__ import annotations

from pathlib import Path

import docx

from . import config


def read_resume_text(path: str | Path) -> str:
    """Extract all text (paragraphs + tables) from a .docx resume."""
    path = config.abspath(path)
    if not path.exists():
        raise FileNotFoundError(f"Resume not found at {path}")
    doc = docx.Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def load_profile(cfg: dict | None = None) -> dict:
    """Bundle resume text + applicant details + search prefs for the pipeline."""
    cfg = cfg or config.load_config()
    resume_path = cfg.get("resume_path", "data/profile/master_resume.docx")
    resume_text = read_resume_text(resume_path)
    return {
        "resume_text": resume_text,
        "applicant": cfg.get("applicant", {}),
        "search": cfg.get("search", {}),
        "resume_path": resume_path,
    }
